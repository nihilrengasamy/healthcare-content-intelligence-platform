"""Build a submission-ready APA-style Word report for the Cotiviti project."""

from __future__ import annotations

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


REPORT_TITLE = "AI-Powered Healthcare Content Intelligence for Payment Integrity at Cotiviti"
AUTHOR_NAME = "Nihil Rengasamy"
SUBTITLE = "Cotiviti Intern Assessment"
DATE_LINE = "June 30, 2026"


BODY_PARAGRAPHS = [
    (
        "Artificial intelligence is becoming more valuable in healthcare operations, especially in "
        "work that depends on reading large amounts of policy and reimbursement content. Billing "
        "rules, coverage policies, clinical guidelines, provider contracts, and regulatory "
        "documents are often stored as lengthy PDFs that are difficult to review quickly. My "
        "project focuses on healthcare content intelligence, which means using AI to turn those "
        "unstructured documents into structured, searchable, and explainable information that teams "
        "can actually use."
    ),
    (
        "This topic fits Cotiviti well because the company operates in healthcare analytics, "
        "payment accuracy, and payment integrity. Those functions rely on correctly interpreting "
        "coverage criteria, coding requirements, prior authorization rules, and documentation "
        "expectations. When analysts must manually read every policy from beginning to end, the "
        "work becomes slow, repetitive, and inconsistent. A content intelligence platform offers a "
        "clear proof of concept for reducing that burden while still keeping human reviewers in "
        "control of important decisions."
    ),
    (
        "Several trends make this topic timely. Healthcare organizations are moving beyond basic AI "
        "experimentation and are looking for practical workflow use cases such as summarization, "
        "search, and decision support. Retrieval-augmented generation is especially relevant because "
        "it grounds answers in source documents instead of relying only on model memory. At the same "
        "time, explainable AI is becoming more important in regulated industries because business "
        "users need to understand why a recommendation was made and what evidence supports it."
    ),
    (
        "The opportunity for Cotiviti is strong. A platform like the one demonstrated in my project "
        "can classify uploaded documents, generate structured summaries, compare old and new policy "
        "versions, answer policy questions through RAG, extract candidate rules, derive feature "
        "sets, and support claim-oriented decisions with transparent explanations. That could improve "
        "speed, reduce repetitive manual work, and make policy interpretation more consistent across "
        "teams. It also creates an audit trail, which is especially valuable when recommendations "
        "need to be justified."
    ),
    (
        "There are also meaningful risks. The biggest threat is hallucination or unsupported output. "
        "If a model gives a confident answer that is not actually supported by the source policy, it "
        "could create compliance or operational problems. A second risk is poor extraction quality. "
        "If the system incorrectly identifies CPT codes, diagnosis codes, exclusions, or prior "
        "authorization requirements, downstream outputs become less reliable. Adoption risk matters "
        "too, because even a strong model will struggle if users do not trust the reasoning."
    ),
    (
        "Because of these risks, I believe Cotiviti should approach this topic as a human-in-the-loop "
        "investment rather than a fully autonomous decision engine. Cotiviti could start by building "
        "an internal analyst copilot for healthcare policy intelligence and then deploy it first in "
        "narrow, high-value areas such as imaging policies, prior authorization workflows, and "
        "reimbursement edits. That phased approach would reduce risk, make impact easier to measure, "
        "and strengthen user trust. Overall, AI-powered healthcare content intelligence is a strong "
        "strategic area for Cotiviti to explore because it aligns with the company’s domain strengths "
        "while addressing a real operational need."
    ),
]


REFERENCES = [
    "Cotiviti. (n.d.-a). Healthcare claim payment accuracy solutions. https://www.cotiviti.com/solutions/payment-accuracy",
    "Cotiviti. (n.d.-b). Healthcare analytics company. https://www.cotiviti.com/",
    "Cotiviti. (2024). The payment integrity guide for health plans. https://info.cotiviti.com/hubfs/assets/white_paper/Cotiviti-Guide-PaymentIntegrity.pdf",
    "Food and Drug Administration. (2024). Artificial intelligence in software as a medical device. https://www.fda.gov/medical-devices/software-medical-device-samd/artificial-intelligence-software-medical-device",
    "McKinsey & Company. (2025). Generative AI in healthcare: Adoption matures as agentic AI emerges. https://www.mckinsey.com/industries/healthcare/our-insights/generative-ai-in-healthcare-current-trends-and-future-outlook",
    "National Institute of Standards and Technology. (2023). AI risk management framework (AI RMF 1.0). https://www.nist.gov/itl/ai-risk-management-framework",
    "Office of the National Coordinator for Health Information Technology. (2024). HTI-1 final rule. https://www.healthit.gov/topic/laws-regulation-and-policy/health-data-technology-and-interoperability-hti-1-final-rule",
]


def set_font(run, size=12, bold=False, italic=False, name="Times New Roman") -> None:
    """Apply font settings to a run."""
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def apply_page_number(section) -> None:
    """Reserve the section header for optional pagination."""
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def style_document(document: Document) -> None:
    """Apply base page and normal style settings."""
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.line_spacing = 2
    normal_style.paragraph_format.space_after = Pt(0)
    normal_style.paragraph_format.space_before = Pt(0)

    for style_name in ["Title", "Subtitle", "Heading 1", "Heading 2"]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")


def add_title_page(document: Document) -> None:
    """Add APA-like title page."""
    section = document.sections[0]
    apply_page_number(section)

    for _ in range(8):
        document.add_paragraph()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = 2
    title_run = title.add_run(REPORT_TITLE)
    set_font(title_run, size=12, bold=True)

    for line in [AUTHOR_NAME, SUBTITLE, DATE_LINE]:
        para = document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.line_spacing = 2
        run = para.add_run(line)
        set_font(run, size=12)


def add_body(document: Document) -> None:
    """Add the report body."""
    section = document.add_section(WD_SECTION.NEW_PAGE)
    apply_page_number(section)

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.line_spacing = 2
    run = heading.add_run(REPORT_TITLE)
    set_font(run, size=12, bold=True)

    for paragraph_text in BODY_PARAGRAPHS:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Inches(0.5)
        paragraph.paragraph_format.line_spacing = 2
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(paragraph_text)
        set_font(run, size=12)


def add_references(document: Document) -> None:
    """Add references page."""
    section = document.add_section(WD_SECTION.NEW_PAGE)
    apply_page_number(section)

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.line_spacing = 2
    run = heading.add_run("References")
    set_font(run, size=12, bold=True)

    for reference in REFERENCES:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.5)
        paragraph.paragraph_format.line_spacing = 2
        run = paragraph.add_run(reference)
        set_font(run, size=12)


def build_document(output_path: str) -> None:
    """Build the Word report document."""
    document = Document()
    style_document(document)
    add_title_page(document)
    add_body(document)
    add_references(document)
    document.save(output_path)


if __name__ == "__main__":
    build_document("Cotiviti_AI_Healthcare_Content_Intelligence_Report.docx")
