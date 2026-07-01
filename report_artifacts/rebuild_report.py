from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = r"C:\Users\Nihil Rengasamy\Documents\Codex\2026-06-27\hi\outputs\healthcare-content-intelligence\Cotiviti_AI_Healthcare_Content_Intelligence_Report_Formatted.docx"

def set_font(run, name="Calibri", size=11, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def add_page_number(paragraph):
    run = paragraph.add_run()
    set_font(run, size=10)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def style_paragraph(paragraph, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=0.5, before=0, after=6, line=1.5):
    pf = paragraph.paragraph_format
    pf.alignment = align
    pf.first_line_indent = Inches(first_indent) if first_indent else None
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line

def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    style_paragraph(p)
    run = p.add_run(text)
    set_font(run)

def add_heading(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = None
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.0
    run = p.add_run(text)
    set_font(run, bold=True, size=11)

def add_reference(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Inches(-0.5)
    pf.left_indent = Inches(0.5)
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.5
    run = p.add_run(text)
    set_font(run)

doc = Document()
sec = doc.sections[0]
for attr in ['top_margin','bottom_margin','left_margin','right_margin']:
    setattr(sec, attr, Inches(1))
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.header_distance = Inches(0.5)
sec.footer_distance = Inches(0.5)
header_p = sec.header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_page_number(header_p)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
p.paragraph_format.line_spacing = 1.0
run = p.add_run("AI-Powered Healthcare Content Intelligence for Payment Integrity at Cotiviti")
set_font(run, bold=True, size=16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
p.paragraph_format.line_spacing = 1.0
run = p.add_run("Nihil Rengasamy | Cotiviti Intern Assessment | June 30, 2026")
set_font(run, size=10)

sections = [
    ("Introduction", [
        "Artificial intelligence is becoming more valuable in healthcare operations, especially in work that depends on reading large volumes of policy and reimbursement content. Billing rules, coverage policies, clinical guidelines, provider contracts, and regulatory documents are often stored as long PDF files that take time to review and compare. In this project, healthcare content intelligence refers to using AI to convert those unstructured documents into structured, searchable, and explainable information that analysts can use more effectively.",
        "This topic is highly relevant to Cotiviti because the company works at the intersection of healthcare analytics, payment accuracy, and payment integrity. Those functions depend on correctly interpreting coverage criteria, coding requirements, prior authorization rules, and documentation expectations. When teams must manually read every policy from beginning to end, the process becomes slow, repetitive, and harder to standardize. A content intelligence platform offers a practical proof of concept for reducing that burden while keeping human reviewers in control of final decisions."
    ]),
    ("Emerging Trends", [
        "Healthcare organizations are moving beyond basic AI experimentation and toward practical workflow use cases such as summarization, semantic search, and decision support. Retrieval-augmented generation, or RAG, is especially useful here because it grounds answers in source documents instead of relying only on model memory. Explainable AI is also becoming more important in regulated industries because users need to understand why a recommendation was made and what evidence supports it."
    ]),
    ("Opportunities for Cotiviti", [
        "The opportunity for Cotiviti is strong. A platform like the one demonstrated in this project can classify uploaded documents, generate structured summaries, compare old and new policy versions, answer policy questions through RAG, extract candidate business rules, derive structured feature sets, and support claim-oriented decisions with transparent explanations. These capabilities can improve review speed, reduce repetitive manual effort, and make policy interpretation more consistent across teams.",
        "There is also strategic value in creating a reusable internal intelligence layer for healthcare content. Instead of treating every policy review as a separate manual task, Cotiviti could build a shared platform that supports analysts, engineers, and business teams with the same source-grounded intelligence."
    ]),
    ("Challenges and Threats", [
        "There are meaningful risks as well. The biggest threat is hallucination or unsupported output. If a model gives a confident answer that is not actually supported by the source policy, it could create compliance problems or poor operational decisions. A second risk is extraction quality. If the system incorrectly identifies CPT codes, diagnosis codes, exclusions, or prior authorization requirements, downstream outputs become less reliable.",
        "Adoption risk also matters. Even a technically strong system can struggle if analysts do not trust the reasoning behind the outputs. In a healthcare environment, users need confidence that the system is grounded, transparent, and easy to validate. That is why explainability, confidence scoring, and human review should remain central to any real deployment approach."
    ]),
    ("Strategic Recommendations for Cotiviti", [
        "Cotiviti should approach this topic as a human-in-the-loop investment rather than a fully autonomous decision engine. A practical next step would be to develop an internal analyst copilot for healthcare policy intelligence and pilot it in narrow, high-value areas such as imaging policies, prior authorization workflows, and reimbursement edits. That phased rollout would reduce implementation risk and make business impact easier to measure.",
        "Cotiviti should also invest in the enabling controls around the AI workflow, not just the model outputs themselves. Strong prompt management, explainability, evaluation dashboards, and source-grounded retrieval will be just as important as the language model."
    ]),
    ("Conclusion", [
        "Overall, AI-powered healthcare content intelligence is a strong strategic area for Cotiviti to explore because it aligns with the company's domain strengths while addressing a real operational need. The proof of concept shows how modern AI can turn static healthcare documents into usable summaries, reusable rules, structured features, and transparent decision support. With phased rollout and human oversight, this concept has clear potential to improve efficiency, consistency, and payment integrity outcomes."
    ])
]
for heading, paras in sections:
    add_heading(doc, heading)
    for para in paras:
        add_body_paragraph(doc, para)

p = doc.add_paragraph(); p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.0; run = p.add_run("References"); set_font(run, bold=True, size=11)
refs = [
    "Cotiviti. (n.d.-a). Healthcare claim payment accuracy solutions. https://www.cotiviti.com/solutions/payment-accuracy",
    "Cotiviti. (n.d.-b). Healthcare analytics company. https://www.cotiviti.com/",
    "Cotiviti. (2024). The payment integrity guide for health plans. https://info.cotiviti.com/hubfs/assets/white_paper/Cotiviti-Guide-PaymentIntegrity.pdf",
    "Food and Drug Administration. (2024). Artificial intelligence in software as a medical device. https://www.fda.gov/medical-devices/software-medical-device-samd/artificial-intelligence-software-medical-device",
    "McKinsey & Company. (2025). Generative AI in healthcare: Adoption matures as agentic AI emerges. https://www.mckinsey.com/industries/healthcare/our-insights/generative-ai-in-healthcare-current-trends-and-future-outlook",
    "National Institute of Standards and Technology. (2023). AI risk management framework (AI RMF 1.0). https://www.nist.gov/itl/ai-risk-management-framework",
    "Office of the National Coordinator for Health Information Technology. (2024). Health data, technology, and interoperability: HTI-1 final rule. https://www.healthit.gov/topic/laws-regulation-and-policy/health-data-technology-and-interoperability-hti-1-final-rule"
]
for ref in refs:
    add_reference(doc, ref)

doc.save(OUT)
print(OUT)

