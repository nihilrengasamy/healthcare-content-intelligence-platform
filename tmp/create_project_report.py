"""Generate a two-page project report for the Healthcare Content Intelligence Platform."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(
    r"C:\Users\Nihil Rengasamy\Documents\Codex\2026-06-27\hi\outputs\healthcare-content-intelligence"
)
OUTPUT_PATH = PROJECT_ROOT / "output" / "Healthcare_Content_Intelligence_Project_Report.docx"


def set_cell_shading(cell, fill: str) -> None:
    """Apply background shading to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_page_margins(document: Document) -> None:
    """Apply consistent page margins."""
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


def set_default_styles(document: Document) -> None:
    """Configure reusable document styles."""
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)

    for style_name, size, color in [
        ("Title", 22, RGBColor(24, 52, 92)),
        ("Heading 1", 14, RGBColor(24, 52, 92)),
        ("Heading 2", 11.5, RGBColor(42, 91, 140)),
    ]:
        style = document.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = color


def add_title(document: Document) -> None:
    """Add the report title block."""
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Healthcare Content Intelligence Platform")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(24, 52, 92)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subrun = subtitle.add_run("Two-Page Project Report for Interview Preparation")
    subrun.italic = True
    subrun.font.size = Pt(11)
    subrun.font.color.rgb = RGBColor(90, 98, 115)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metarun = meta.add_run("Project Type: Enterprise AI + Healthcare NLP + Streamlit")
    metarun.font.size = Pt(10)
    metarun.font.color.rgb = RGBColor(70, 70, 70)


def add_overview_callout(document: Document) -> None:
    """Add a compact executive callout."""
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.8)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EAF2FB")
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(
        "This project is an AI-powered Healthcare Content Intelligence Platform built "
        "to analyze policy documents, extract healthcare rules and features, support "
        "policy question answering through RAG, and demonstrate explainable decision support "
        "for claims-oriented workflows."
    )
    run.font.size = Pt(10)
    run.bold = True


def add_bullets(document: Document, items: list[str]) -> None:
    """Add bullet items with compact spacing."""
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.05
        paragraph.add_run(item)


def add_page_one(document: Document) -> None:
    """Add page one content."""
    document.add_heading("1. Project Objective", level=1)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.add_run(
        "The platform was designed as an end-to-end proof of concept for healthcare "
        "content management. Its goal is to help analysts and payment integrity teams "
        "work faster with complex policy PDFs such as billing and coding policies, "
        "clinical practice guidelines, payer-provider contracts, coverage policies, "
        "and regulatory documents."
    )

    document.add_heading("2. Business Problem Being Solved", level=1)
    add_bullets(
        document,
        [
            "Healthcare policy documents are long, unstructured, and difficult to compare manually.",
            "Analysts need fast access to medical necessity criteria, authorization rules, coding rules, and exclusions.",
            "Traditional review is time-consuming and can miss version changes, contract rules, or documentation requirements.",
            "An interview-ready platform should show document intelligence, RAG, rule extraction, ML signals, and explainability in one workflow.",
        ],
    )

    document.add_heading("3. End-to-End Workflow", level=1)
    workflow = document.add_paragraph()
    workflow.paragraph_format.space_after = Pt(4)
    workflow.add_run(
        "Repository -> PDF Loader -> Document Classification -> Summarization -> "
        "Version Comparison -> Embeddings -> FAISS Vector Store -> Policy Chat (RAG) "
        "-> Rule Extraction -> Feature Extraction -> Rule Engine -> ML Prediction "
        "-> Claim Decision -> Explainability -> Evaluation -> Analytics Dashboard"
    ).bold = True

    document.add_heading("4. Core Technical Architecture", level=1)
    add_bullets(
        document,
        [
            "Frontend: Streamlit multi-page enterprise dashboard for document upload, AI workflows, and analytics.",
            "Document ingestion: PyMuPDF extracts page-level text and metadata from PDF policies.",
            "Chunking: LangChain RecursiveCharacterTextSplitter creates chunks for embeddings and retrieval.",
            "Retrieval layer: sentence-transformers embeddings + FAISS vector index enable semantic policy search.",
            "LLM layer: Groq/OpenAI-backed modules produce summaries, RAG answers, and structured rule extraction.",
        ],
    )

    document.add_heading("5. Main Modules in the Project", level=1)
    modules_table = document.add_table(rows=1, cols=2)
    modules_table.style = "Table Grid"
    headers = modules_table.rows[0].cells
    headers[0].text = "Module"
    headers[1].text = "Purpose"
    for cell in headers:
        set_cell_shading(cell, "D9E7F7")
        for run in cell.paragraphs[0].runs:
            run.bold = True

    module_rows = [
        ("pdf_loader.py", "Loads healthcare PDFs, extracts pages, metadata, and text chunks."),
        ("document_classifier.py", "Classifies uploaded documents into billing, clinical, contract, coverage, or regulatory types."),
        ("summarizer.py", "Generates structured healthcare policy summaries for analyst review."),
        ("compare_versions.py", "Compares old and new policy versions and highlights meaningful changes."),
        ("embeddings.py + vector_store.py", "Create document embeddings and a FAISS index for semantic retrieval."),
        ("rag.py", "Answers policy questions using source-grounded Retrieval-Augmented Generation."),
    ]
    for module_name, purpose in module_rows:
        row = modules_table.add_row().cells
        row[0].text = module_name
        row[1].text = purpose


def add_page_break(document: Document) -> None:
    """Insert a page break."""
    document.add_page_break()


def add_page_two(document: Document) -> None:
    """Add page two content."""
    document.add_heading("6. Advanced Intelligence Layers", level=1)
    add_bullets(
        document,
        [
            "rule_extractor.py converts policy text into machine-readable JSON business rules.",
            "feature_extractor.py converts healthcare text into structured features such as CPT, ICD, therapy weeks, prior authorization, and contract terms.",
            "rule_engine.py evaluates extracted rules against claim or patient features.",
            "ml_model.py adds predictive signals such as approval probability, fraud risk, and medical necessity score.",
            "claim_decision.py combines rule-based and ML-based outputs into approve, deny, or manual review recommendations.",
            "explainability.py and evaluation.py improve auditability, groundedness, and interview strength by showing why the system made a decision.",
        ],
    )

    document.add_heading("7. Technology Stack", level=1)
    tech = document.add_paragraph()
    tech.paragraph_format.space_after = Pt(4)
    tech.add_run(
        "Python, Streamlit, LangChain, Groq/OpenAI-compatible LLM APIs, FAISS, PyMuPDF, "
        "sentence-transformers, Pandas, Scikit-learn, Plotly, and Pydantic."
    )

    document.add_heading("8. Why This Project Is Strong for Interview Discussion", level=1)
    add_bullets(
        document,
        [
            "It is not a single-script demo; it demonstrates modular architecture and separation of concerns.",
            "It covers multiple AI patterns: summarization, RAG, rule extraction, structured NLP, deterministic rules, ML scoring, and explainability.",
            "It shows practical healthcare domain thinking, including medical necessity, prior authorization, CPT/ICD coding, and coverage policies.",
            "It includes a usable frontend, so the project can be demonstrated live instead of explained only through code.",
        ],
    )

    document.add_heading("9. Current Strengths and Honest Gaps", level=1)
    strengths_table = document.add_table(rows=1, cols=2)
    strengths_table.style = "Table Grid"
    head = strengths_table.rows[0].cells
    head[0].text = "Strengths"
    head[1].text = "Current Gaps / Next Improvements"
    for cell in head:
        set_cell_shading(cell, "D9E7F7")
        for run in cell.paragraphs[0].runs:
            run.bold = True

    rows = [
        (
            "End-to-end workflow from document upload to analytics dashboard.",
            "Production deployment concerns such as authentication, PHI controls, and audit logging would need hardening.",
        ),
        (
            "Clear module boundaries and test coverage across major components.",
            "Some answers still depend on retrieval quality and prompt tuning, especially for ambiguous policy questions.",
        ),
        (
            "Strong interview value because it combines healthcare + AI + product demo.",
            "Real-world rollout would need curated data, stronger governance, and validation by domain experts.",
        ),
    ]
    for left, right in rows:
        row = strengths_table.add_row().cells
        row[0].text = left
        row[1].text = right

    document.add_heading("10. Interview Closing Summary", level=1)
    closing = document.add_paragraph()
    closing.paragraph_format.space_after = Pt(0)
    closing.add_run(
        "In summary, this project demonstrates how AI can convert unstructured healthcare content into "
        "searchable knowledge, analyst-friendly summaries, executable rules, structured features, and "
        "decision support signals. The strongest way to present it in an interview is to emphasize that "
        "the platform is both architecturally modular and operationally demonstrable: it shows document "
        "intelligence, RAG, rules, ML, explainability, and analytics inside one coherent healthcare workflow."
    )


def add_footer(document: Document) -> None:
    """Add page footer to the first section."""
    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Healthcare Content Intelligence Platform Project Report")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)


def main() -> None:
    """Create the report document."""
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    set_page_margins(document)
    set_default_styles(document)
    add_title(document)
    add_overview_callout(document)
    add_page_one(document)
    add_page_break(document)
    add_page_two(document)
    add_footer(document)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
